from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.dml.color import ColorFormat
from docx.oxml.ns import qn

import datetime, os, sys, time


class Handle_lxb():

    def __init__(self, oversea_group, staff_datas, template_address, target_address, desk_address):
        self.template_path = template_address + "立项表_模板.docx"
        self.document = Document(self.template_path)
        self.paragraphs = self.document.paragraphs
        self.tables = self.document.tables
        self.group_members = oversea_group['group_members']
        self.group = oversea_group
        self.staff = staff_datas
        self.target_path = \
            target_address + \
            " - ".join([
                time.strftime("%Y%m%d", time.localtime()),
                "立项表",
                self.group_members[0],
                self.group['destination'],
                str(len(self.group_members)) + "人.docx"
            ])
        self.user_path = str(
            desk_address + "/立项表/" + " - ".join(
                [time.strftime("%Y%m%d", time.localtime()), "立项表", self.group_members[0],
                 self.group['destination'], str(len(self.group_members)) + "人.docx"])).replace("\\", "/")

        self.title_style = self.document.styles.add_style('title_style', 1)
        self.title_style.font.size = Pt(15)
        self.title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.title_style.font.name = '宋体'
        self.title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.title_style.font.bold = True

        self.user_comment_style = self.document.styles.add_style('user_comment_style', 2)
        self.user_comment_style.font.size = Pt(14)
        # self.user_comment_style.characters_format.alignment = WD_ALIGN_PARAGRAPH.CENTER   #字符型居中格式命令方式待定
        self.user_comment_style.font.name = '宋体'
        self.user_comment_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.user_comment_style.font.bold = False

        self.table_style = self.document.styles.add_style('table_style', 3)
        self.table_style.font.size = Pt(14)
        self.table_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.table_style.font.name = '宋体'
        self.table_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.table_style.font.bold = False

    def fill_lxb(self):

        if self.group['group_leader_position'] < 1:
            self.paragraphs[2].text = "所属企业主要领导出国、赴港澳立项审查表"
        elif self.group['group_leader_position'] < 2:
            self.paragraphs[2].text = "所属企业非正职领导干部出国、赴港澳立项审查表"
        else:
            self.paragraphs[2].text = "员工出国、赴港澳立项审查表"
        self.paragraphs[2].style = self.title_style

        self.tables[0].cell(0, 1).text = "哈电国际赴" + self.group['destination'] + "工作团组"
        self.tables[0].cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        self.tables[0].cell(0, 10).text = self.group_members[0]
        self.tables[0].cell(0, 10).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.tables[0].cell(1, 1).text = self.group['job_detail']
        self.tables[0].cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        self.tables[0].cell(1, 10).text = "/".join(
            [self.group['departure_date_year'], self.group['departure_date_month'], self.group['departure_date_day']])
        self.tables[0].cell(1, 10).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if int(self.group['time']) < 180:
            self.tables[0].cell(2, 5).text = self.group['destination'] + '（多次往返，本次停留' + str(self.group['time']) + '天）'
            self.tables[0].cell(2, 10).text = '365'
        else:
            self.tables[0].cell(2, 5).text = self.group['destination']
            self.tables[0].cell(2, 10).text = str(self.group['time'])
        self.tables[0].cell(2, 5).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.tables[0].cell(2, 10).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.tables[0].cell(3, 2).text = self.group['inviter']
        self.tables[0].cell(3, 2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        self.tables[0].cell(4, 2).text = self.group['cost_center']
        self.tables[0].cell(4, 2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.tables[0].cell(4, 12).text = self.group['combined_in_heg']
        self.tables[0].cell(4, 12).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_row_no = 6
        for group_member in self.group_members:
            for staff_data in self.staff:
                if group_member == staff_data['name']:
                    self.tables[0].cell(cell_row_no, 0).text = staff_data['name']
                    self.tables[0].cell(cell_row_no, 1).text = staff_data['gender']
                    self.tables[0].cell(cell_row_no, 2).text = str(staff_data['age'])
                    self.tables[0].cell(cell_row_no, 3).text = staff_data['company']
                    self.tables[0].cell(cell_row_no, 8).text = staff_data['position']
                    self.tables[0].cell(cell_row_no, 0).paragraphs[
                        0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.tables[0].cell(cell_row_no, 1).paragraphs[
                        0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.tables[0].cell(cell_row_no, 2).paragraphs[
                        0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.tables[0].cell(cell_row_no, 3).paragraphs[
                        0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.tables[0].cell(cell_row_no, 8).paragraphs[
                        0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    break
                else:
                    continue
            cell_row_no += 1
        if len(self.group_members) >= 6:
            self.tables[0].cell(cell_row_no, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif len(self.group_members) == 5:
            self.tables[0].cell(cell_row_no, 0).text = '此行空白'
        else:
            self.tables[0].cell(cell_row_no, 0).text = '以下空白'
        self.tables[0].cell(cell_row_no, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_row_no = 6
        #self.tables[0].style = self.table_style
        self.document.save(self.target_path)
        self.document.save(self.user_path)
