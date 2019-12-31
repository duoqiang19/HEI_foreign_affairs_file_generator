from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor, ElementProxy, Emu, lazyproperty, Length, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.dml.color import ColorFormat
from docx.oxml.ns import qn
# from docx.tabstops import TabStops

import datetime, os, sys, time


class Handle_qswj():

    def __init__(self, oversea_group, staff_datas, template_address, target_address, desk_address):
        self.template_path = template_address + "请示文件_模板.docx"
        self.document = Document(self.template_path)
        self.paragraphs = self.document.paragraphs
        self.tables = self.document.tables
        self.group_members = oversea_group['group_members']
        self.group = oversea_group
        self.staff = staff_datas
        self.target_path = \
            target_address + \
            " - ".join([
                "哈电国际外事〔2020〕" + str(self.group['foreign_affair_doc_no']) + "号",
                self.group['group_leader'],
                self.group['destination'],
                str(self.group['group_members_no']) + "人.docx"
            ])
        self.user_path = str(
            desk_address + "/请示文件/" + " - ".join(
                ["哈电国际外事〔2020〕" + str(self.group['foreign_affair_doc_no']) + "号", self.group['group_leader'],
                 self.group['destination'], str(self.group['group_members_no']) + "人.docx"])).replace("\\", "/")

        self.title_style = self.document.styles.add_style('title_style', 1)
        self.title_style.font.size = Pt(22)
        self.title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.title_style.font.name = '新宋体'
        self.title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '新宋体')
        self.title_style.font.bold = True

        self.normal_style = self.document.styles.add_style('normal_style', 1)
        self.normal_style.font.size = Pt(16)
        self.normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.normal_style._element.get_or_add_pPr().spacing_after = Pt(0)
        self.normal_style._element.get_or_add_pPr().spacing_before = Pt(0)
        self.normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        self.normal_style._element.get_or_add_pPr().line_spacing_rule = Pt(18)
        self.normal_style.font.name = '仿宋_GB2312'
        self.normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        self.normal_style.font.bold = False

        self.table_character_style = self.document.styles.add_style('table_character_style', 2)
        self.table_character_style.font.size = Pt(16)
        self.table_character_style.font.name = '仿宋_GB2312'
        self.table_character_style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        self.table_character_style.font.bold = False

    def fill_qswj(self):

        self.paragraphs[8].text = \
            "关于拟派" + \
            self.group['group_leader'] + \
            "等" + \
            str(self.group['group_members_no']) + \
            "人赴" + \
            self.group['destination'] + \
            "工作的"
        self.paragraphs[8].style = self.title_style

        self.paragraphs[12].text = \
            "应" + \
            self.group['inviter'] + \
            "的邀请，我公司拟派" + \
            self.group['group_leader'] + \
            "等" + \
            str(self.group['group_members_no']) + \
            "人于" + \
            self.group['departure_date_year'] + \
            '年' + \
            self.group['departure_date_month'] + \
            '月' + \
            self.group['departure_date_day'] + \
            "日赴" + \
            self.group['destination'] + \
            "执行参加" + \
            self.group['job_detail'] + \
            "等工作任务。"
        self.paragraphs[12].style = self.normal_style

        self.tables[0].cell(0, 0).text = '哈电国际外事〔2020〕' + str(self.group['foreign_affair_doc_no']) + '号'
        self.tables[0].cell(0, 0).paragraphs[0].style = self.normal_style
        self.tables[0].cell(0, 2).text = ''
        self.tables[0].cell(0, 2).paragraphs[0].style = self.normal_style

        for group_member in self.group_members:
            for staff_data in self.staff:
                if group_member == staff_data['name']:
                    # self.paragraphs[13].paragraph_format.next_paragraph_style.delete()
                    self.paragraphs[14].insert_paragraph_before( \
                        staff_data['name'] + \
                        '\t\t' + \
                        staff_data['gender'] + \
                        '\t\t' + \
                        str(staff_data['age']) + \
                        '岁\t\t哈电国际\t\t' + \
                        staff_data['position'])
                    # next_paragraph_style
                    break
                else:
                    continue

        self.paragraphs[15].text = \
            self.group['departure_date_year'] + \
            '年' + \
            self.group['departure_date_month'] + \
            '月' + \
            self.group['departure_date_day'] + \
            "日，本次停留" + \
            str(self.group['time']) + \
            '天'
        self.paragraphs[15].style = self.normal_style

        self.paragraphs[17].text = "→".join(self.group['tour_rountine'])
        self.paragraphs[17].style = self.normal_style

        self.paragraphs[23].text = \
            '                               ' + \
            time.strftime("%Y", time.localtime()) + \
            '年' + \
            time.strftime("%m", time.localtime()) + \
            '月' + \
            time.strftime("%d", time.localtime()) + \
            '日'
        self.paragraphs[23].style = self.normal_style

        self.tables[1].cell(0, 1).text = \
            time.strftime("%Y", time.localtime()) + \
            '年' + \
            time.strftime("%m", time.localtime()) + \
            '月' + \
            time.strftime("%d", time.localtime()) + \
            '日印发'
        self.tables[1].cell(0, 1).paragraphs[0].style = self.normal_style
        self.tables[1].cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        self.document.save(self.target_path)
        self.document.save(self.user_path)
