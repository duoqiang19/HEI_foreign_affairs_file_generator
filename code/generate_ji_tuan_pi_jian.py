from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.dml.color import ColorFormat
from docx.oxml.ns import qn

import datetime, os, sys, time


class Handle_jtpj():

    def __init__(self, oversea_group, staff_datas, template_address, target_address, desk_address):

        self.template_path = template_address + "任务批件_模板.docx"
        self.document = Document(self.template_path)
        self.paragraphs = self.document.paragraphs
        self.tables = self.document.tables
        self.group_members = oversea_group['group_members']
        self.group = oversea_group
        self.staff = staff_datas
        self.target_path = \
            target_address + \
            " - ".join([
                '哈电集团出任字（ 2020HP ）' + str(self.group['he_permission_no']) + '号',
                self.group_members[0],
                str(len(self.group_members)) + "人",
                self.group['destination'] + '.docx'
            ])
        self.user_path = str(desk_address + "/集团批件/" + " - ".join(
            ['哈电集团出任字（ 2020HP ）' + str(self.group['he_permission_no']) + '号', self.group_members[0],
             str(len(self.group_members)) + "人", self.group['destination'] + '.docx'])).replace("\\", "/")

        self.he_permission_no_style = self.document.styles.add_style('he_permission_no_style', 1)
        self.he_permission_no_style.font.size = Pt(12)
        self.he_permission_no_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.he_permission_no_style.font.name = 'Times New Roman'
        self.he_permission_no_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.he_permission_no_style.font.bold = False

        self.table_comment_style = self.document.styles.add_style('table_comment_style', 1)
        self.table_comment_style.font.size = Pt(14)
        self.table_comment_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.table_comment_style.font.name = '宋体'
        self.table_comment_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.table_comment_style.font.bold = False

        self.date_style = self.document.styles.add_style('date_style', 1)
        self.date_style.font.size = Pt(10)
        self.date_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.date_style.font.name = '宋体'
        self.date_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.date_style.font.bold = False

    def fill_jtpj(self):

        self.paragraphs[3].text = '哈电集团出任字（ 2020HP ）' + str(self.group['he_permission_no']) + '号'
        self.paragraphs[3].style = self.he_permission_no_style

        self.tables[0].cell(0, 1).text = \
            "哈电国际" + \
            self.group_members[0] + \
            "等" + \
            str(len(self.group_members)) + \
            "人赴" + \
            self.group['destination'] + \
            "工作团组"
        self.tables[0].cell(0, 1).paragraphs[0].style = self.table_comment_style
        self.tables[0].cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        self.tables[0].cell(0, 9).text = self.group_members[0]
        self.tables[0].cell(0, 9).paragraphs[0].style = self.table_comment_style
        self.tables[0].cell(0, 9).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.tables[0].cell(1, 1).text = self.group['job_detail']
        self.tables[0].cell(1, 1).paragraphs[0].style = self.table_comment_style
        self.tables[0].cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if int(self.group['time']) < 180:
            self.tables[0].cell(2, 5).text = self.group['destination'] + '（多次往返，本次停留' + str(self.group['time']) + '天）'
            self.tables[0].cell(2, 10).text = '365'
        else:
            self.tables[0].cell(2, 5).text = self.group['destination']
            self.tables[0].cell(2, 10).text = str(self.group['time'])
        self.tables[0].cell(2, 5).paragraphs[0].style = self.table_comment_style
        self.tables[0].cell(2, 10).paragraphs[0].style = self.table_comment_style
        self.tables[0].cell(2, 5).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.tables[0].cell(2, 10).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.tables[0].cell(4, 1).text = self.group['inviter']
        self.tables[0].cell(4, 1).paragraphs[0].style = self.table_comment_style
        self.tables[0].cell(4, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.tables[0].cell(5, 1).text = self.group['cost_center']
        self.tables[0].cell(5, 1).paragraphs[0].style = self.table_comment_style
        self.tables[0].cell(5, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_row_no = 7
        for group_member in self.group_members:
            for staff_data in self.staff:
                if group_member == staff_data['name']:
                    self.tables[0].cell(cell_row_no, 0).text = staff_data['name']
                    self.tables[0].cell(cell_row_no, 2).text = staff_data['gender']
                    self.tables[0].cell(cell_row_no, 3).text = str(staff_data['age'])
                    self.tables[0].cell(cell_row_no, 4).text = staff_data['company']
                    self.tables[0].cell(cell_row_no, 10).text = staff_data['position']
                    self.tables[0].cell(cell_row_no, 0).paragraphs[0].style = self.table_comment_style
                    self.tables[0].cell(cell_row_no, 2).paragraphs[0].style = self.table_comment_style
                    self.tables[0].cell(cell_row_no, 3).paragraphs[0].style = self.table_comment_style
                    self.tables[0].cell(cell_row_no, 4).paragraphs[0].style = self.table_comment_style
                    self.tables[0].cell(cell_row_no, 10).paragraphs[0].style = self.table_comment_style
                    self.tables[0].cell(cell_row_no, 0).paragraphs[
                        0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.tables[0].cell(cell_row_no, 2).paragraphs[
                        0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.tables[0].cell(cell_row_no, 3).paragraphs[
                        0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.tables[0].cell(cell_row_no, 4).paragraphs[
                        0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.tables[0].cell(cell_row_no, 10).paragraphs[
                        0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break
                else:
                    continue
            cell_row_no += 1
        if len(self.group_members) >= 7:
            self.tables[0].cell(cell_row_no, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif len(self.group_members) == 6:
            self.tables[0].cell(cell_row_no, 0).text = '此行空白'
        else:
            self.tables[0].cell(cell_row_no, 0).text = '以下空白'
        self.tables[0].cell(cell_row_no, 0).paragraphs[0].style = self.table_comment_style
        self.tables[0].cell(cell_row_no, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.paragraphs[4].text = \
            '附  注：拟出访日期：' + \
            '.'.join(
                [self.group['departure_date_year'],
                 self.group['departure_date_month'],
                 self.group['departure_date_day']
                 ]) + \
            '（停留时间包含抵离境当日）            （公    章）'
        self.paragraphs[4].style = self.date_style

        self.document.save(self.target_path)
        self.document.save(self.user_path)
